"""Тесты импорта ДС из Word (.docx)."""
from __future__ import annotations

import io
from decimal import Decimal
from io import BytesIO

from docx import Document

from app.modules.reference.amendment_docx_import import parse_amendment_docx
from app.modules.reference.models import TariffRule

_TARIFF_ROWS = [
    ["№", "Наименование", "Ед.", "Тариф"],
    [
        "1",
        "Хранение на площади 9435 m2 за сутки",
        "м2",
        "24,00 руб.",
    ],
    [
        "2",
        "Механизированная обработка (вход и выход), м³",
        "м3",
        "180,00 руб.",
    ],
]


def _make_docx(rows: list[list[str]], paragraphs: list[str] | None = None, cols: int = 4) -> bytes:
    doc = Document()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    if rows:
        table = doc.add_table(rows=len(rows), cols=cols)
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                table.rows[i].cells[j].text = val
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_amendment_docx_table_tariffs():
    docx = _make_docx(_TARIFF_ROWS, paragraphs=["Дополнительное соглашение № 5 от 01.01.2026"])
    parsed = parse_amendment_docx(docx, "ДС 5.docx")
    codes = {t.billing_line_code for t in parsed.tariffs}
    assert "storage_area_fixed" in codes
    assert "mechanized_m3" in codes
    mech = next(t for t in parsed.tariffs if t.billing_line_code == "mechanized_m3")
    assert mech.rate == Decimal("180")


def test_amendment_import_creates_tariff_rules(auth_client, client, app):
    auth_client("admin@test.local", "admin")
    contracts = client.get("/api/reference/contracts").json["items"]
    contract_id = contracts[0]["id"]

    docx = _make_docx(_TARIFF_ROWS, paragraphs=["Дополнительное соглашение № 5 от 01.01.2026"])

    data = {
        "contract_id": str(contract_id),
        "file": (io.BytesIO(docx), "test_ds.docx"),
    }
    resp = client.post(
        "/api/reference/amendments/import",
        data=data,
        content_type="multipart/form-data",
    )
    assert resp.status_code == 201
    body = resp.json
    assert body["tariffs_created"] >= 2
    amendment_id = body["amendment"]["id"]

    tariffs_resp = client.get("/api/reference/tariff_rules")
    assert tariffs_resp.status_code == 200
    linked = [t for t in tariffs_resp.json["items"] if t["amendment_id"] == amendment_id]
    assert len(linked) >= 2

    lookups = client.get("/api/reference/lookups").json
    assert any(a["id"] == amendment_id for a in lookups["amendments"])

    with app.app_context():
        rules = TariffRule.query.filter_by(amendment_id=amendment_id).all()
        assert len(rules) >= 2
        assert all(r.contract_id == contract_id for r in rules)
