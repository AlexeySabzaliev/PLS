"""Парсинг допсоглашений ОХ из DOCX: таблицы произвольной структуры и текст."""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import date
from decimal import Decimal, InvalidOperation
from io import BytesIO

from docx import Document

from app.modules.uss.services.tariff_codes import (
    formula_for_code,
    infer_billing_line_code,
    unit_code_for_billing_line,
)

MONTHS = {
    "январ": 1, "феврал": 2, "март": 3, "апрел": 4, "ма": 5, "июн": 6,
    "июл": 7, "август": 8, "сентябр": 9, "октябр": 10, "ноябр": 11, "декабр": 12,
}

UNIT_CODES = {
    "м2": "m2", "м²": "m2", "m2": "m2", "кв.м": "m2", "кв. м": "m2", "квм": "m2",
    "м3": "m3", "м³": "m3", "m3": "m3", "куб": "m3",
    "пакет": "vehicle", "авто": "vehicle", "машин": "vehicle", "тс": "vehicle",
    "шт": "pcs", "штук": "pcs", "ед": "pcs", "ед.": "pcs",
    "час": "hour", "чел": "hour", "чел.ч": "hour", "чел.-час": "hour",
}

HEADER_NAME = ("наимен", "услуг", "вид работ", "описан", "позици", "работ")
HEADER_UNIT = ("ед.", "ед ", "единиц", "изм")
HEADER_RATE = ("тариф", "ставк", "цена", "стоим", "руб", "сумм", "размер", "оплат")
HEADER_NUM = ("№", "номер", "п/п", "п.п", "n ")

UNPRICED_MARKERS = (
    "по согласованию",
    "по договоренности",
    "не согласован",
    "не согласована",
    "цена не согласован",
    "стоимость не согласован",
    "без указания стоимости",
    "без указания цены",
    "определяется отдельно",
    "определяется по согласованию",
)

NAME_STOP_MARKERS = (
    "в данный тариф входит",
    "в данный тариф входят",
    "тариф включает",
    "в стоимость входит",
    "в стоимость тарифа входит",
)


@dataclass
class ParsedTariff:
    billing_line_code: str
    name: str
    rate: Decimal
    unit_code: str | None = None
    formula: str = "rate_times_qty"
    sort_order: int = 0


@dataclass
class ParsedParameter:
    param_type: str
    numeric_value: Decimal


@dataclass
class ParsedAmendment:
    number: str | None = None
    amendment_type: str = "rate_change"
    effective_from: date | None = None
    effective_to: date | None = None
    signed_date: date | None = None
    description: str | None = None
    source_document: str | None = None
    tariffs: list[ParsedTariff] = field(default_factory=list)
    parameters: list[ParsedParameter] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def _to_decimal(value: str) -> Decimal | None:
    cleaned = value.strip().replace(" ", "").replace(",", ".")
    cleaned = re.sub(r"[^\d.]", "", cleaned)
    if not cleaned:
        return None
    try:
        return Decimal(cleaned)
    except InvalidOperation:
        return None


def _extract_rates(text: str) -> list[Decimal]:
    found: list[Decimal] = []
    for m in re.finditer(r"(\d{1,6}[,.]\d{2})\s*(?:руб|₽)?", text, re.IGNORECASE):
        d = _to_decimal(m.group(1))
        if d is not None:
            found.append(d)
    return found


def _parse_ru_date(text: str) -> date | None:
    m = re.search(r"(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{4})", text)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        try:
            return date(y, mo, d)
        except ValueError:
            return None
    m = re.search(
        r"(?:с\s+)?(\d{1,2})\s+([а-яё]+)\s+(\d{4})",
        text.lower(),
    )
    if m:
        day = int(m.group(1))
        month_word = m.group(2)
        year = int(m.group(3))
        month = next((n for k, n in MONTHS.items() if month_word.startswith(k)), None)
        if month:
            try:
                return date(year, month, day)
            except ValueError:
                return None
    return None


def _parse_all_dates(text: str) -> list[date]:
    found: list[date] = []
    for m in re.finditer(r"(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{4})", text):
        try:
            found.append(date(int(m.group(3)), int(m.group(2)), int(m.group(1))))
        except ValueError:
            continue
    for m in re.finditer(r"(\d{1,2})\s+([а-яё]+)\s+(\d{4})", text.lower()):
        month = next((n for k, n in MONTHS.items() if m.group(2).startswith(k)), None)
        if month:
            try:
                found.append(date(int(m.group(3)), month, int(m.group(1))))
            except ValueError:
                continue
    return found


def _parse_effective_from(paragraphs: list, all_text: str) -> date | None:
    period_from, _ = _parse_period_range(all_text)
    if period_from:
        return period_from
    priority: list[date] = []
    for para in paragraphs:
        t = para.text
        if not t.strip():
            continue
        low = t.lower()
        if any(k in low for k in (
            "вступает в силу",
            "действует с",
            "вносят изменения",
            "изменения вносятся",
            "приложение",
            "срок действия",
        )):
            d = _parse_ru_date(t)
            if d:
                priority.append(d)
    if priority:
        return max(priority)
    dates = _parse_all_dates(all_text)
    return max(dates) if dates else None


def _parse_period_range(all_text: str) -> tuple[date | None, date | None]:
    patterns = (
        r"действует\s+с\s+"
        r"(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{4}|\d{1,2}\s+[а-яё]+\s+\d{4})\s+"
        r"по\s+"
        r"(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{4}|\d{1,2}\s+[а-яё]+\s+\d{4})",
        r"с\s+(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{4}|\d{1,2}\s+[а-яё]+\s+\d{4})\s+"
        r"по\s+"
        r"(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{4}|\d{1,2}\s+[а-яё]+\s+\d{4})",
    )
    for pat in patterns:
        m = re.search(pat, all_text, re.IGNORECASE)
        if not m:
            continue
        start = _parse_ru_date(m.group(1))
        end = _parse_ru_date(m.group(2))
        if start and end and end >= start:
            return start, end
    return None, None


def _parse_effective_to(paragraphs: list, all_text: str) -> date | None:
    period_from, period_to = _parse_period_range(all_text)
    if period_to and period_to != period_from:
        return period_to
    priority: list[date] = []
    for para in paragraphs:
        t = para.text
        if not t.strip():
            continue
        low = t.lower()
        if any(k in low for k in (
            "действует по",
            "сроком до",
            "окончан",
            "по состоянию на",
        )):
            d = _parse_ru_date(t)
            if d:
                priority.append(d)
    if priority:
        return max(priority)
    return None


def _normalize_effective_dates(
    effective_from: date | None,
    effective_to: date | None,
) -> tuple[date | None, date | None]:
    if effective_from and effective_to and effective_to <= effective_from:
        return effective_from, None
    return effective_from, effective_to


def _guess_unit(text: str) -> str | None:
    low = text.lower()
    for key, code in sorted(UNIT_CODES.items(), key=lambda x: -len(x[0])):
        if key in low:
            return code
    return None


def _row_label(cell: str) -> str | None:
    m = re.match(r"^(\d+(?:\.\d+)?)", cell.strip())
    return m.group(1) if m else None


def _row_number(cell: str) -> int | None:
    label = _row_label(cell)
    if not label:
        return None
    return int(label.split(".")[0])


def _is_storage_extra_name(low: str) -> bool:
    if "дополнительн" not in low:
        return False
    return any(k in low for k in ("площад", "объем", "объём", "м2", "м²"))


def _extract_storage_area_params(text: str) -> list[ParsedParameter]:
    params: list[ParsedParameter] = []
    seen: set[str] = set()
    for m in re.finditer(r"(\d[\d\s]{1,6})\s*м[²2]", text, re.IGNORECASE):
        val = int(re.sub(r"\s+", "", m.group(1)))
        ptype = "fixed_storage_m2days"
        if ptype not in seen:
            params.append(ParsedParameter(ptype, Decimal(val)))
            seen.add(ptype)
    return params


def _pick_tariff_rate(text: str, rates: list[Decimal]) -> Decimal | None:
    if not rates:
        return None
    for pattern in (
        r"(?:ставк[аи]|тариф\w*|из\s+расчёта|стоимост\w*)\s*[^\d]{0,40}?(\d{1,4}[,.]\d{2})\s*руб",
        r"(\d{1,4}[,.]\d{2})\s*руб",
    ):
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            picked = _to_decimal(m.group(1))
            if picked is not None:
                return picked
    plausible = [r for r in rates if r >= Decimal("5")]
    if plausible:
        return max(plausible)
    return rates[-1]


def _header_match(cell: str, keywords: tuple[str, ...]) -> bool:
    low = cell.lower().strip()
    return any(k in low for k in keywords)


def _detect_table_columns(header: list[str]) -> dict[str, int | None]:
    cols: dict[str, int | None] = {"num": None, "name": None, "unit": None, "rate": None}
    for i, cell in enumerate(header):
        low = cell.lower().strip()
        if not low:
            continue
        if cols["name"] is None and _header_match(low, HEADER_NAME):
            cols["name"] = i
        elif cols["unit"] is None and _header_match(low, HEADER_UNIT):
            cols["unit"] = i
        elif cols["rate"] is None and _header_match(low, HEADER_RATE):
            cols["rate"] = i
        elif cols["num"] is None and (_header_match(low, HEADER_NUM) or low in ("#", "n")):
            cols["num"] = i
    return cols


def _looks_like_header_row(cells: list[str]) -> bool:
    joined = " ".join(cells).lower()
    has_name = _header_match(joined, HEADER_NAME)
    has_rate = _header_match(joined, HEADER_RATE)
    return has_name or has_rate


def _compact_name(text: str) -> str:
    return re.sub(r"\s+", "", (text or "").lower())


def _is_manual_handling_name(low: str) -> bool:
    compact = _compact_name(low)
    return "ручн" in compact or ("руч" in compact and "обработ" in compact)


def _has_volume_in_name(low: str) -> bool:
    compact = _compact_name(low)
    return "м3" in compact or "м³" in low or "короб" in compact or "1м3" in compact


def _clean_tariff_name(raw: str) -> str:
    text = (raw or "").replace("\r", "\n").strip()
    low = text.lower()
    for marker in NAME_STOP_MARKERS:
        idx = low.find(marker)
        if idx > 0:
            text = text[:idx].strip()
            low = text.lower()
    if "•" in text:
        text = text.split("•", 1)[0].strip()
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if lines:
        text = lines[0]
    return text.strip()


def _cell_title(cell) -> str:
    """Заголовок услуги: жирный текст в ячейке или первая строка без описания."""
    bold_parts: list[str] = []
    for para in cell.paragraphs:
        for run in para.runs:
            if run.bold and (run.text or "").strip():
                bold_parts.append(run.text.strip())
    if bold_parts:
        return _clean_tariff_name(" ".join(bold_parts))
    return _clean_tariff_name(cell.text)


def _is_unpriced_row(*texts: str) -> bool:
    combined = " ".join(t for t in texts if t).lower()
    return any(marker in combined for marker in UNPRICED_MARKERS)


def _display_name(raw: str, fallback: str) -> str:
    text = _clean_tariff_name(raw or "")
    return text[:255] if len(text) >= 3 else fallback


def _finalize_tariff(t: ParsedTariff) -> ParsedTariff:
    code = infer_billing_line_code(t.name, t.billing_line_code)
    unit = t.unit_code or unit_code_for_billing_line(code) or _guess_unit(t.name)
    formula = formula_for_code(code) if code in {
        "storage_area_fixed", "storage_area_extra", "storage_area",
        "manual_m3", "extra_manual_m3", "mechanized_m3", "vehicle_docs",
        "extra_vehicle_docs", "repack_units", "overtime_m3",
        "inventory_hours", "elco_drain_hours",
    } else t.formula
    return ParsedTariff(
        billing_line_code=code,
        name=t.name.strip()[:255],
        rate=t.rate,
        unit_code=unit,
        formula=formula,
        sort_order=t.sort_order,
    )


def _heuristic_row_tariff(cells: list[str], sort_order: int) -> ParsedTariff | None:
    texts = [c.strip() for c in cells if c and c.strip()]
    if len(texts) < 2:
        return None

    if _is_unpriced_row(*texts):
        return None

    rate: Decimal | None = None
    rate_idx: int | None = None
    for i, t in enumerate(texts):
        picked = _pick_tariff_rate(t, _extract_rates(t))
        if picked is not None:
            rate = picked
            rate_idx = i
            if re.search(r"руб|₽", t, re.I) or len(t) < 40:
                break

    if rate is None:
        return None

    unit_idx: int | None = None
    unit_text = ""
    for i, t in enumerate(texts):
        if i == rate_idx:
            continue
        if _guess_unit(t) and len(t) < 30:
            unit_text = t
            unit_idx = i
            break

    name_parts: list[str] = []
    for i, t in enumerate(texts):
        if i in (rate_idx, unit_idx):
            continue
        if re.fullmatch(r"[\d.\s]+", t):
            continue
        if _pick_tariff_rate(t, _extract_rates(t)) and len(t) < 50:
            continue
        if len(t) >= 2:
            name_parts.append(t)
    name = _clean_tariff_name(" ".join(name_parts).strip())
    if len(name) < 3:
        return None

    code = infer_billing_line_code(name, None)
    unit = _guess_unit(unit_text or name) or unit_code_for_billing_line(code)
    return _finalize_tariff(ParsedTariff(code, name, rate, unit, formula_for_code(code), sort_order))


def _parse_amendment_number(all_text: str, filename: str) -> str | None:
    patterns = (
        r"(?:дополнительн\w+\s+соглашени\w+|дс)\s*[№#]?\s*(\d+(?:/\d{4})?)",
        r"дс\s*[-–]?\s*(\d+(?:/\d{4})?)",
        r"№\s*(\d+)\s*(?:к\s+договору)?",
    )
    for pat in patterns:
        m = re.search(pat, all_text, re.IGNORECASE)
        if m:
            num = m.group(1)
            if "/" not in num and re.search(r"20\d{2}", all_text):
                year_m = re.search(r"(20\d{2})", all_text)
                if year_m:
                    return f"ДС-{num}/{year_m.group(1)}"
            return f"ДС-{num}" if not num.upper().startswith("ДС") else num
    m_file = re.search(r"дс\s*(\d+)", filename, re.I) if filename else None
    if m_file:
        return f"ДС-{m_file.group(1)}"
    return None


def _map_table_row_legacy(
    row_label: str | None,
    row_num: int | None,
    name: str,
    unit_text: str,
    rates: list[Decimal],
) -> tuple[list[ParsedTariff], list[ParsedParameter]]:
    tariffs: list[ParsedTariff] = []
    params: list[ParsedParameter] = []
    low = name.lower()

    if row_label == "1.2" or (_is_storage_extra_name(low) and row_label != "1"):
        if rates:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "storage_area_extra",
                _display_name(name, "Площадь хранения, дополнительный объём, м²"),
                rates[0], "m2", "rate_times_days_times_qty", 12,
            )))
        return tariffs, params

    if row_num == 1 or ("9435" in name and "м2" in low):
        if rates:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "storage_area_fixed",
                _display_name(name, "Площадь хранения, фиксированный объём, м²"),
                rates[0], "m2", "rate_times_days_times_qty", 11,
            )))
        params.extend(_extract_storage_area_params(name))
        return tariffs, params

    if "дополнительн" in low and "ручн" in low and "обработ" in low:
        if rates:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "extra_manual_m3",
                _display_name(name, "Дополнительная ручная обработка, м³"),
                rates[0], "m3", sort_order=21,
            )))
        return tariffs, params

    if "паллет" in low:
        if rates and not _is_unpriced_row(name, unit_text):
            tariffs.append(_finalize_tariff(ParsedTariff(
                "custom_pallet",
                _display_name(name, "Дополнительная обработка при паллетировании грузов"),
                rates[0], "pcs", sort_order=55,
            )))
        return tariffs, params

    if _is_manual_handling_name(low) and _has_volume_in_name(low):
        if rates:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "manual_m3",
                _display_name(name, "Ручная обработка (вход и выход), м³"),
                rates[0], "m3", sort_order=20,
            )))
        return tariffs, params

    if ("механиз" in low or "механизир" in _compact_name(low)) and _has_volume_in_name(low):
        if rates:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "mechanized_m3",
                _display_name(name, "Механизированная обработка (вход и выход), м³"),
                rates[0], "m3", sort_order=30,
            )))
        return tariffs, params

    if row_num == 2 or ("погруз" in low and "разгруз" in low and "ручн" not in low):
        if len(rates) >= 2:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "mechanized_m3",
                _display_name(name, "Механизированная обработка (вход и выход), м³"),
                rates[0], "m3", sort_order=30,
            )))
            tariffs.append(_finalize_tariff(ParsedTariff(
                "manual_m3", "Ручная обработка (вход и выход), м³",
                rates[1], "m3", sort_order=20,
            )))
        elif len(rates) == 1:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "mechanized_m3",
                _display_name(name, "Механизированная обработка (вход и выход), м³"),
                rates[0], "m3", sort_order=30,
            )))
        return tariffs, params

    if row_num == 4 or (
        ("пакет" in low and "документ" in low)
        and "ручн" not in low
        and "механиз" not in low
        and "паллет" not in low
    ):
        if rates:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "vehicle_docs",
                _display_name(name, "Количество пакетов документов (машин)"),
                rates[0], "vehicle", sort_order=40,
            )))
        return tariffs, params

    if row_num == 4 or ("механиз" in low and "м3" in low):
        return tariffs, params

    if row_num == 5 or "переупак" in low:
        if rates:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "repack_units", "Переупаковка (единиц приборов)",
                rates[0], "pcs", sort_order=50,
            )))
        return tariffs, params

    if "сверхур" in low:
        if rates and not _is_unpriced_row(name, unit_text):
            tariffs.append(_finalize_tariff(ParsedTariff(
                "overtime_m3", "Сверхурочная обработка авто (вход и выход), м³",
                rates[0], "m3", sort_order=60,
            )))
        return tariffs, params

    if "инвентариз" in low:
        if rates and not _is_unpriced_row(name, unit_text):
            tariffs.append(_finalize_tariff(ParsedTariff(
                "inventory_hours",
                _display_name(name, "Дополнительная инвентаризация"),
                rates[0], "hour", sort_order=70,
            )))
        return tariffs, params

    if "слив" in low and ("elco" in low or "элко" in low):
        if rates:
            tariffs.append(_finalize_tariff(ParsedTariff(
                "elco_drain_hours", "Слив технологической жидкости ELCO",
                rates[0], "hour", sort_order=80,
            )))
        return tariffs, params

    if rates and row_num:
        tariffs.append(_finalize_tariff(ParsedTariff(
            f"custom_row_{row_num}",
            _display_name(name, f"Услуга {row_num}"),
            rates[0],
            _guess_unit(unit_text or name),
            sort_order=100 + (row_num or 0),
        )))
    return tariffs, params


def _parse_table_row(
    cells: list[str],
    columns: dict[str, int | None],
    row_index: int,
    cell_objs: list | None = None,
) -> tuple[list[ParsedTariff], list[ParsedParameter]]:
    if not any(c.strip() for c in cells):
        return [], []

    def cell(key: str, default: str = "") -> str:
        idx = columns.get(key)
        if idx is None or idx >= len(cells):
            return default
        if key == "name" and cell_objs and idx < len(cell_objs):
            return _cell_title(cell_objs[idx])
        return cells[idx].strip()

    name = cell("name")
    unit_text = cell("unit")
    rate_text = cell("rate")
    num_text = cell("num", cells[0] if cells else "")

    if not rate_text:
        for c in cells:
            if _extract_rates(c):
                rate_text = c
                break

    rates = _extract_rates(rate_text)
    if not rates and unit_text:
        rates = _extract_rates(unit_text)

    if _is_unpriced_row(name, unit_text, rate_text, *cells):
        return [], []

    row_label = _row_label(num_text)
    row_num = _row_number(num_text or (cells[0] if cells else ""))

    if not name:
        extra_name_parts: list[str] = []
        for i, c in enumerate(cells):
            if columns.get("name") == i:
                continue
            if i == columns.get("rate") or i == columns.get("unit"):
                continue
            if len(c.strip()) > 0 and not re.fullmatch(r"[\d.\s]+", c.strip()):
                extra_name_parts.append(c.strip())
        name = _clean_tariff_name(" ".join(extra_name_parts).strip())

    legacy_tariffs, legacy_params = _map_table_row_legacy(
        row_label, row_num, name, unit_text, rates,
    )
    if legacy_tariffs:
        return legacy_tariffs, legacy_params

    if name and rates:
        rate = _pick_tariff_rate(rate_text or name, rates)
        if rate is not None:
            t = _heuristic_row_tariff(cells, 10 + row_index)
            if t:
                return [t], _extract_storage_area_params(name)
            code = infer_billing_line_code(name, None)
            unit = _guess_unit(unit_text or name) or unit_code_for_billing_line(code)
            return [_finalize_tariff(ParsedTariff(
                code, name, rate, unit, formula_for_code(code), 10 + row_index,
            ))], _extract_storage_area_params(name)

    t = _heuristic_row_tariff(cells, 10 + row_index)
    if t:
        return [t], _extract_storage_area_params(t.name)
    return [], []


def _parse_table(table, start_order: int = 0) -> tuple[list[ParsedTariff], list[ParsedParameter]]:
    if not table.rows:
        return [], []
    tariffs: list[ParsedTariff] = []
    params: list[ParsedParameter] = []
    header_cells = [c.text.strip() for c in table.rows[0].cells]
    columns = _detect_table_columns(header_cells)
    data_start = 1 if _looks_like_header_row(header_cells) else 0

    if columns["name"] is None and columns["rate"] is None and len(table.rows[0].cells) >= 3:
        columns = {"num": 0, "name": 1, "unit": 2, "rate": 3 if len(header_cells) > 3 else 2}

    for ri, row in enumerate(table.rows[data_start:]):
        cells = [c.text.strip() for c in row.cells]
        cell_objs = list(row.cells)
        if not any(cells):
            continue
        row_tariffs, row_params = _parse_table_row(cells, columns, start_order + ri, cell_objs)
        tariffs.extend(row_tariffs)
        params.extend(row_params)

    return tariffs, params


def _find_extra_storage_tariff(text: str) -> ParsedTariff | None:
    low = text.lower()
    if not _is_storage_extra_name(low):
        return None
    rates = _extract_rates(text)
    rate = _pick_tariff_rate(text, rates)
    if rate is None:
        return None
    return _finalize_tariff(ParsedTariff(
        "storage_area_extra",
        "Площадь хранения, дополнительный объём, м²",
        rate, "m2", "rate_times_days_times_qty", 12,
    ))


def _parse_text_tariffs(text: str, base_order: int = 500) -> list[ParsedTariff]:
    tariffs: list[ParsedTariff] = []
    seen_rates: set[tuple[str, str]] = set()

    line_patterns = (
        re.compile(
            r"(?:^|\n)\s*(?:\d+[\.\):]|\-|\•)\s*"
            r"(?P<name>.+?)"
            r"\s*(?P<rate>\d{1,6}[,.]\d{2})\s*(?:руб|₽)"
            r"(?:\s*(?:за|/)\s*(?P<unit>[^\n,.;]{1,40}))?",
            re.IGNORECASE | re.MULTILINE,
        ),
        re.compile(
            r"(?P<name>[А-Яа-яЁё0-9][^\n]{4,}?)"
            r"\s*[-–—:]\s*"
            r"(?P<rate>\d{1,6}[,.]\d{2})\s*(?:руб|₽)"
            r"(?:\s*(?:за|/)\s*(?P<unit>[^\n,.;]{1,40}))?",
            re.IGNORECASE,
        ),
        re.compile(
            r"(?:тариф|ставк)\w*\s+(?:в\s+размере\s+)?(?P<rate>\d{1,6}[,.]\d{2})\s*(?:руб|₽)"
            r".{0,80}?(?:за|/)\s*(?P<unit>[^\n,.;]{1,25})",
            re.IGNORECASE,
        ),
    )

    order = base_order
    for pat in line_patterns:
        for m in pat.finditer(text):
            name = _clean_tariff_name((m.groupdict().get("name") or "").strip(" .;:-–—"))
            rate_raw = m.groupdict().get("rate")
            if not name or not rate_raw:
                if rate_raw and not name:
                    ctx_start = max(0, m.start() - 120)
                    name = _clean_tariff_name(text[ctx_start:m.start()].strip().split("\n")[-1].strip(" .;:-–—"))
                if not name or len(name) < 8:
                    continue
            if _is_unpriced_row(name, m.group(0)):
                continue
            rate = _to_decimal(rate_raw or "")
            if rate is None:
                continue
            unit_raw = (m.groupdict().get("unit") or "").strip()
            key = (name[:80].lower(), str(rate))
            if key in seen_rates:
                continue
            seen_rates.add(key)
            code = infer_billing_line_code(name, None)
            unit = _guess_unit(unit_raw or name) or unit_code_for_billing_line(code)
            tariffs.append(_finalize_tariff(ParsedTariff(
                code, name[:255], rate, unit, formula_for_code(code), order,
            )))
            order += 1

    for para_text in text.split("\n"):
        extra = _find_extra_storage_tariff(para_text)
        if extra:
            key = (extra.billing_line_code, str(extra.rate))
            if key not in seen_rates:
                seen_rates.add(key)
                tariffs.append(extra)

    return tariffs


def _merge_tariffs(tariff_lists: list[list[ParsedTariff]]) -> list[ParsedTariff]:
    by_code: dict[str, ParsedTariff] = {}
    for lst in tariff_lists:
        for t in lst:
            by_code[t.billing_line_code] = t
    return sorted(by_code.values(), key=lambda x: (x.sort_order, x.name))


def parse_amendment_docx(file_bytes: bytes, filename: str = "") -> ParsedAmendment:
    doc = Document(BytesIO(file_bytes))
    result = ParsedAmendment(source_document=filename or None)
    paragraphs = list(doc.paragraphs)
    all_text = "\n".join(p.text for p in paragraphs if p.text.strip())

    result.number = _parse_amendment_number(all_text, filename)
    result.effective_from = _parse_effective_from(paragraphs, all_text)
    result.effective_to = _parse_effective_to(paragraphs, all_text)
    result.effective_from, result.effective_to = _normalize_effective_dates(
        result.effective_from, result.effective_to,
    )

    signed = _parse_ru_date(filename) if filename else None
    if signed:
        result.signed_date = signed

    all_tariffs: list[list[ParsedTariff]] = []
    all_params: list[ParsedParameter] = []
    order_base = 0
    for table in doc.tables:
        t_list, p_list = _parse_table(table, order_base)
        if t_list:
            all_tariffs.append(t_list)
            order_base += len(t_list)
        all_params.extend(p_list)

    text_tariffs = _parse_text_tariffs(all_text)
    if text_tariffs:
        all_tariffs.append(text_tariffs)

    result.tariffs = _merge_tariffs(all_tariffs)

    by_param: dict[str, ParsedParameter] = {p.param_type: p for p in all_params}
    for p in _extract_storage_area_params(all_text):
        by_param.setdefault(p.param_type, p)
    result.parameters = list(by_param.values())

    if "storage_area_extra" not in {t.billing_line_code for t in result.tariffs}:
        for para in paragraphs:
            extra = _find_extra_storage_tariff(para.text)
            if extra:
                result.tariffs.append(extra)
                break

    if not result.tariffs:
        result.warnings.append("tariffs_not_found")
    if not result.effective_from:
        result.warnings.append("effective_from_not_found")

    result.description = result.description or "Импорт из DOCX"
    return result


def parsed_to_dict(parsed: ParsedAmendment) -> dict:
    return {
        "number": parsed.number,
        "amendment_type": parsed.amendment_type,
        "effective_from": parsed.effective_from.isoformat() if parsed.effective_from else None,
        "effective_to": parsed.effective_to.isoformat() if parsed.effective_to else None,
        "signed_date": parsed.signed_date.isoformat() if parsed.signed_date else None,
        "description": parsed.description,
        "source_document": parsed.source_document,
        "warnings": parsed.warnings,
        "tariffs": [
            {
                "billing_line_code": t.billing_line_code,
                "name": t.name,
                "rate": float(t.rate),
                "unit_code": t.unit_code,
                "formula": t.formula,
                "sort_order": t.sort_order,
            }
            for t in parsed.tariffs
        ],
        "parameters": [
            {"param_type": p.param_type, "numeric_value": float(p.numeric_value)}
            for p in parsed.parameters
        ],
    }
